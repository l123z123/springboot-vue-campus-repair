"""
Apply humanized text back to the docx file.
Preserves front matter (paragraphs 0-58) untouched.
Reads humanized JSON files and writes changes to a new docx.
"""
import docx
from docx.shared import Pt, Cm
import json
import os
from copy import deepcopy

BASE = r"D:\Office-software\project\springboot+vue"
DOCX_IN = os.path.join(BASE, "thesis_formatted_v2.docx")
DOCX_OUT = os.path.join(BASE, "thesis_humanized.docx")

HUMANIZED_FILES = [
    "docs/论文写作/extracted_abstract_cn_humanized.json",
    "docs/论文写作/extracted_abstract_en_humanized.json",
    "docs/论文写作/extracted_ch1_xulun_humanized.json",
    "docs/论文写作/extracted_ch2_tech_humanized.json",
    "docs/论文写作/extracted_ch3_analysis_humanized.json",
    "docs/论文写作/extracted_ch4_design_humanized.json",
    "docs/论文写作/extracted_ch5_impl_humanized.json",
    "docs/论文写作/extracted_ch6_test_humanized.json",
    "docs/论文写作/extracted_ch7_conclusion_humanized.json",
    "docs/论文写作/extracted_acknowledgments_humanized.json",
]

def clear_paragraph(p):
    """Clear all runs from a paragraph while preserving the paragraph element."""
    for run in p.runs:
        run.text = ''

def set_paragraph_text(p, text):
    """Set paragraph text, preserving formatting from existing runs."""
    if not text:
        clear_paragraph(p)
        return

    if p.runs:
        # Preserve first run formatting, put all text there
        first_run = p.runs[0]
        # Clear all other runs
        for run in p.runs[1:]:
            run.text = ''
        first_run.text = text
    else:
        # Add a new run
        p.add_run(text)

def main():
    print(f"Loading original docx: {DOCX_IN}")
    doc = docx.Document(DOCX_IN)

    # Load all humanized paragraphs into a lookup dict
    humanized = {}
    for fname in HUMANIZED_FILES:
        fpath = os.path.join(BASE, fname)
        if os.path.exists(fpath):
            with open(fpath, 'r', encoding='utf-8') as f:
                items = json.load(f)
            for item in items:
                idx = item['index']
                humanized[idx] = item['text']
            print(f"  Loaded {len(items)} paragraphs from {os.path.basename(fname)}")
        else:
            print(f"  WARNING: {fpath} not found, skipping")

    print(f"Total humanized paragraphs: {len(humanized)}")

    # Verify no front-matter paragraphs (0-58) are in the humanized set
    front_matter_touched = [i for i in humanized if i < 59]
    if front_matter_touched:
        print(f"ERROR: Front matter paragraphs would be modified: {front_matter_touched}")
        print("Aborting to protect front matter.")
        return
    print("Front matter (paragraphs 0-58) protected: no changes will be made.")

    # Apply changes
    changed = 0
    skipped = 0
    for idx in sorted(humanized.keys()):
        if idx >= len(doc.paragraphs):
            print(f"  WARNING: paragraph {idx} out of range (max {len(doc.paragraphs)-1}), skipping")
            skipped += 1
            continue

        new_text = humanized[idx]
        p = doc.paragraphs[idx]
        old_text = p.text

        if old_text != new_text:
            set_paragraph_text(p, new_text)
            changed += 1

    print(f"Changed {changed} paragraphs, skipped {skipped}")
    print(f"Saving to: {DOCX_OUT}")
    doc.save(DOCX_OUT)

    # Verify output
    doc2 = docx.Document(DOCX_OUT)
    print(f"Output docx has {len(doc2.paragraphs)} paragraphs")
    print("Done!")

if __name__ == '__main__':
    main()
